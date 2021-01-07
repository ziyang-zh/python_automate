import docx

with open('guests.txt','r') as txtFile:
	guests=txtFile.read().split('\n')
print(guests)

doc=docx.Document()
for guest in guests:
	doc.add_paragraph('It would be a pleasure to have the company of')
	doc.add_paragraph(guest)
	doc.add_paragraph('at 11010 memory Lane on the Evening of')
	doc.add_paragraph('April 1st')
	doc.add_paragraph("at 7 o'clock")
	doc.paragraphs[-1].runs[-1].add_break(docx.text.run.WD_BREAK.PAGE)
doc.save('guests.docx')