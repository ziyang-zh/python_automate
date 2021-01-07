import docx

doc=docx.Document('demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)
print(doc.paragraphs[1].runs[4].text)

def getText(filename):
	doc=docx.Document(filename)
	fullText=[]
	for para in doc.paragraphs:
		fullText.append(para.text)
	return '\n'.join(fullText)

print(getText('demo.docx'))

doc=docx.Document('demo.docx')
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
doc.paragraphs[0].style='Normal'
print(doc.paragraphs[1].text)
print((doc.paragraphs[1].runs[0].text,
		doc.paragraphs[1].runs[1].text,
		doc.paragraphs[1].runs[2].text,
		doc.paragraphs[1].runs[3].text,
		doc.paragraphs[1].runs[4].text,))
doc.paragraphs[1].runs[0].style='QuoteChar'
doc.paragraphs[1].runs[1].underline=True
doc.paragraphs[1].runs[3].underline=True
doc.save('restyled.docx')

doc=docx.Document()
doc.add_paragraph('Hello world!','Title')
doc.save('helloworld.docx')
paraObj1=doc.add_paragraph('This is a second paragraph.')
paraObj2=doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('multipleParagraphs.docx')

doc=docx.Document()
doc.add_heading('Header 0',0)
doc.add_heading('Header 1',1)
doc.add_heading('Header 2',2)
doc.add_heading('Header 3',3)
doc.add_heading('Header 4',4)
doc.save('headings.docx')

doc=docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.text.run.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')

doc.add_picture('zophie.png',width=docx.shared.Inches(1),height=docx.shared.Cm(4))
